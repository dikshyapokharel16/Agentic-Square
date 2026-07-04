"""Sync the chat script between messages.json and a Word document.

Only plain dialogue (dates, system notes, messages, image/stage changes,
level cards) is editable from Word. Polls, file shares, and the visitor
reply-prompt carry data that has no simple Word representation (poll
options, reaction/vote wiring, prompt copy, tap-choice/keyword-bucket
wiring) and stay defined in messages.json only — they show up in the doc
as a single non-editable "[LOCKED: ...]" marker line so editors can see
the story's shape without being able to touch that content.

Usage:
    python chat_docx.py export   # messages.json -> chat-script.docx
    python chat_docx.py sync     # chat-script.docx -> messages.json
"""

import argparse
import colorsys
import hashlib
import json
import re
from pathlib import Path

from docx import Document

LOCKED_TYPES = {"poll", "file", "userPrompt"}

# Pinned so Marktplatz's default look (no role tag) keeps a stable brand
# color (matches --accent-terracotta in styles.css) instead of being
# hash-generated like community members. Most Marktplatz lines carry an
# explicit [Role] tag instead (see color_for/ROLE_STYLES in main.js) —
# this is only the fallback for a roleless line.
BRAND_COLORS = {"Marktplatz": ("#BF5468", "#fbe7ea", "L")}

HEADER_LINES = [
    "Agentic Square — Chat Script",
    "Edit this document freely: reword messages, add new ones, remove ones, "
    "reorder them. When you're done, ask Claude to \"sync the chat script\" "
    "(or run: python tools/chat_docx.py sync) and it'll update messages.json "
    "to match.",
    "",
    "Each message is one block, separated by a line that says exactly \"---\". "
    "The first line of a block is TYPE, sender (if any), and time. Everything "
    "after is the message content.",
    "",
    "Types you can use:",
    "  DATE hh:mm            -> a date divider, e.g. \"Monday\"",
    "  SYSTEM hh:mm          -> a small centered system note",
    "  EVENT hh:mm           -> like SYSTEM but highlighted",
    "  MSG Sender hh:mm      -> a normal chat message",
    "  MSG You hh:mm (me)    -> a message sent by the visitor (right-aligned)",
    "  IMAGE Sender hh:mm    -> shares a picture — ALSO ADVANCES which stage/",
    "                           3D model the AR panel shows, in the order",
    "                           these IMAGE blocks appear top to bottom",
    "  LEVEL n hh:mm         -> a pop-up card next to the chat that pauses it",
    "                           until tapped. First line below is the title,",
    "                           second line is the explanation text.",
    "  LEVEL n hh:mm (auto:4) -> same, but the 4 entries right after this one",
    "                           play out automatically while the card is still",
    "                           up, instead of also waiting on a tap.",
    "",
    "A LEVEL card can also be pinned to a specific already-revealed message's",
    "on-screen height instead of a fixed spot (entry.anchorIndex) — that's a",
    "position in messages.json, not something Word has a stable way to point",
    "at, so it has no Word representation and won't survive a sync; set/edit",
    "it directly in messages.json.",
    "",
    "For MSG/IMAGE lines where Marktplatz is speaking, you can tag which role",
    "it's speaking as — e.g. \"MSG [Designer] Marktplatz 12:04\" — to set that",
    "message's avatar icon and accent color (see the role table in main.js). ",
    "Leave the tag off for the default brand look.",
    "",
    "Type {{name}} anywhere you want the visitor's own name to appear — it's",
    "replaced automatically once someone enters their name on the kiosk.",
    "",
    "Lines that look like \"[LOCKED: ...]\" are polls, file shares, or the",
    "visitor reply-prompt — not editable from Word. Leave them exactly where",
    "they are; edit that content in messages.json directly instead.",
    "",
    "Colors and avatar initials are handled automatically from each sender's",
    "name — just focus on the words. Renaming a sender gives them a new,",
    "consistent color.",
]

TYPE_RE = re.compile(r"^(DATE|SYSTEM|EVENT|MSG|IMAGE|LEVEL)\s+(.*)$", re.IGNORECASE)
TIME_RE = re.compile(r"(\d{1,2}:\d{2})")
LOCKED_RE = re.compile(r"^\[LOCKED:\s*(\w+)\b")


def entry_summary(entry):
    if entry["type"] == "poll":
        return entry.get("question", "poll")
    if entry["type"] == "file":
        return entry.get("fileName", "file")
    if entry["type"] == "userPrompt":
        return entry.get("promptText", "reply prompt")
    return entry["type"]


def locked_marker(entry):
    return f'[LOCKED: {entry["type"]} — "{entry_summary(entry)}" — edit this one in messages.json instead]'


def hsl_to_hex(h, s, l):
    r, g, b = colorsys.hls_to_rgb(h, l, s)
    return "#{:02X}{:02X}{:02X}".format(round(r * 255), round(g * 255), round(b * 255))


def color_for(sender, cache):
    if sender in BRAND_COLORS:
        return BRAND_COLORS[sender]
    if sender in cache:
        return cache[sender]
    digest = hashlib.sha256(sender.encode("utf-8")).hexdigest()
    hue = int(digest[:8], 16) / 0xFFFFFFFF
    bg = hsl_to_hex(hue, 0.45, 0.40)
    fg = hsl_to_hex(hue, 0.55, 0.93)
    initial = sender.strip()[:1].upper() or "?"
    result = (bg, fg, initial)
    cache[sender] = result
    return result


def export_docx(messages_path, docx_path):
    messages = json.loads(Path(messages_path).read_text(encoding="utf-8"))
    doc = Document()

    for line in HEADER_LINES:
        doc.add_paragraph(line)
    doc.add_paragraph("")

    for entry in messages:
        kind = entry["type"]
        doc.add_paragraph("---")

        if kind in LOCKED_TYPES:
            doc.add_paragraph(locked_marker(entry))
        elif kind in ("date", "system", "event"):
            doc.add_paragraph(f"{kind.upper()} {entry.get('time', '')}")
            doc.add_paragraph(entry.get("text", ""))
        elif kind == "level":
            auto_tag = f" (auto:{entry['autoReveal']})" if entry.get("autoReveal") else ""
            doc.add_paragraph(f"LEVEL {entry.get('level', '')} {entry.get('time', '')}{auto_tag}")
            doc.add_paragraph(entry.get("title", ""))
            doc.add_paragraph(entry.get("text", ""))
        elif kind == "msg":
            is_me = bool(entry.get("isMe"))
            sender = "You" if is_me else entry.get("sender", "")
            role_tag = f"[{entry['role']}] " if entry.get("role") else ""
            suffix = " (me)" if is_me else ""
            doc.add_paragraph(f"MSG {role_tag}{sender} {entry.get('time', '')}{suffix}")
            doc.add_paragraph(entry.get("text", ""))
        elif kind == "image":
            role_tag = f"[{entry['role']}] " if entry.get("role") else ""
            doc.add_paragraph(f"IMAGE {role_tag}{entry.get('sender', '')} {entry.get('time', '')}")
            doc.add_paragraph(entry.get("caption", ""))
        else:
            raise ValueError(f"Don't know how to export entry type {kind!r}")

        doc.add_paragraph("")

    doc.save(docx_path)
    print(f"Exported {messages_path} -> {docx_path} ({len(messages)} entries)")


def read_docx_lines(docx_path):
    return [p.text for p in Document(docx_path).paragraphs]


def split_blocks(lines):
    """Blocks start right after a line that's exactly '---' and run until the
    next one; anything before the first '---' (the header) is ignored."""
    blocks = []
    current = None
    for line in lines:
        if line.strip() == "---":
            if current is not None:
                blocks.append(current)
            current = []
        elif current is not None:
            current.append(line)
    if current is not None:
        blocks.append(current)
    return blocks


ROLE_RE = re.compile(r"^\[([^\]]+)\]\s*(.*)$")
AUTO_REVEAL_RE = re.compile(r"\(auto:(\d+)\)\s*$")


def parse_type_line(line):
    match = TYPE_RE.match(line.strip())
    if not match:
        return None
    kind = match.group(1).upper()
    rest = match.group(2).strip()

    # Only meaningful on MSG/IMAGE (which role Marktplatz is speaking as),
    # but harmless to strip unconditionally — no other kind's sender/level
    # field is ever bracketed.
    role = None
    role_match = ROLE_RE.match(rest)
    if role_match:
        role = role_match.group(1)
        rest = role_match.group(2)

    # Only meaningful on LEVEL — how many entries right after it auto-play
    # while the card is still up (see entry.autoReveal / handleLevelPopup
    # in main.js) — but likewise harmless to strip unconditionally.
    auto_reveal = None
    auto_match = AUTO_REVEAL_RE.search(rest)
    if auto_match:
        auto_reveal = int(auto_match.group(1))
        rest = rest[: auto_match.start()].strip()

    is_me = False
    if rest.endswith("(me)"):
        is_me = True
        rest = rest[: -len("(me)")].strip()

    time_match = TIME_RE.search(rest)
    if not time_match:
        raise ValueError(f"Missing hh:mm time in line: {line!r}")
    time = time_match.group(1)
    sender = rest[: time_match.start()].strip()
    return kind, sender, time, is_me, role, auto_reveal


def reactions_key(entry):
    sender = "{{name}}" if entry.get("isMe") else entry.get("sender", "")
    return (entry["type"], sender, entry.get("text") or entry.get("caption") or "")


def sync_from_docx(docx_path, messages_path, out_path=None):
    blocks = split_blocks(read_docx_lines(docx_path))
    original = json.loads(Path(messages_path).read_text(encoding="utf-8"))

    locked_queues = {t: [e for e in original if e["type"] == t] for t in LOCKED_TYPES}
    locked_idx = {t: 0 for t in LOCKED_TYPES}
    color_cache = {}
    result = []

    # msg/image entries can carry a `reactions` field that has no Word
    # representation (out of scope — see plan). Since dialogue can be freely
    # added/reordered/reworded from Word, position isn't a safe way to match
    # entries back up — instead, match on exact (type, sender, text), so
    # reactions survive for every line that wasn't actually changed, and are
    # only (correctly) dropped for lines that were reworded or are new.
    reactions_by_key = {}
    for entry in original:
        if entry["type"] in ("msg", "image") and entry.get("reactions"):
            reactions_by_key.setdefault(reactions_key(entry), []).append(entry["reactions"])

    for block in blocks:
        non_empty = [(i, l) for i, l in enumerate(block) if l.strip()]
        if not non_empty:
            continue
        first_idx, type_line = non_empty[0]
        type_line = type_line.strip()
        body_lines = [l.strip() for l in block[first_idx + 1 :] if l.strip()]
        body = " ".join(body_lines)

        locked_match = LOCKED_RE.match(type_line)
        if locked_match:
            kind = locked_match.group(1)
            if kind not in locked_queues:
                raise ValueError(f"Unknown locked type {kind!r} in: {type_line!r}")
            queue = locked_queues[kind]
            i = locked_idx[kind]
            if i >= len(queue):
                raise ValueError(
                    f"Found more '{kind}' [LOCKED] markers in the doc than messages.json has "
                    f"{kind} entries ({len(queue)}). Polls/files/prompts can't be added or "
                    f"removed from Word — edit messages.json directly for that."
                )
            result.append(queue[i])
            locked_idx[kind] += 1
            continue

        parsed = parse_type_line(type_line)
        if not parsed:
            raise ValueError(f"Couldn't parse block starting with: {type_line!r}")
        kind, sender, time, is_me, role, auto_reveal = parsed

        if kind in ("DATE", "SYSTEM", "EVENT"):
            result.append({"type": kind.lower(), "time": time, "text": body})
        elif kind == "LEVEL":
            entry = {
                "type": "level",
                "level": int(sender) if sender.isdigit() else sender,
                "time": time,
                "title": body_lines[0] if body_lines else "",
                "text": " ".join(body_lines[1:]),
            }
            if auto_reveal:
                entry["autoReveal"] = auto_reveal
            result.append(entry)
        elif kind == "MSG":
            entry = {"type": "msg", "time": time, "text": body}
            if is_me:
                entry["sender"] = "{{name}}"
                entry["isMe"] = True
            else:
                bg, fg, initial = color_for(sender, color_cache)
                entry.update({"sender": sender, "initial": initial, "bg": bg, "fg": fg})
                if role:
                    entry["role"] = role
            matches = reactions_by_key.get(reactions_key(entry))
            if matches:
                entry["reactions"] = matches.pop(0)
            result.append(entry)
        elif kind == "IMAGE":
            bg, fg, initial = color_for(sender, color_cache)
            entry = {
                "type": "image",
                "sender": sender,
                "time": time,
                "initial": initial,
                "bg": bg,
                "fg": fg,
                "caption": body,
            }
            if role:
                entry["role"] = role
            matches = reactions_by_key.get(reactions_key(entry))
            if matches:
                entry["reactions"] = matches.pop(0)
            result.append(entry)

    for kind, queue in locked_queues.items():
        if locked_idx[kind] != len(queue):
            raise ValueError(
                f"messages.json has {len(queue)} '{kind}' entries but the doc only had "
                f"{locked_idx[kind]} [LOCKED: {kind}] markers — don't delete locked markers "
                f"from Word."
            )

    out_path = out_path or messages_path
    Path(out_path).write_text(json.dumps(result, indent=2, ensure_ascii=False) + "\n", encoding="utf-8")
    print(f"Synced {docx_path} -> {out_path} ({len(result)} entries)")


def main():
    parser = argparse.ArgumentParser(description=__doc__)
    sub = parser.add_subparsers(dest="command", required=True)

    p_export = sub.add_parser("export", help="Regenerate chat-script.docx from messages.json")
    p_export.add_argument("--messages", default="../messages.json")
    p_export.add_argument("--docx", default="../chat-script.docx")

    p_sync = sub.add_parser("sync", help="Regenerate messages.json from chat-script.docx")
    p_sync.add_argument("--docx", default="../chat-script.docx")
    p_sync.add_argument("--messages", default="../messages.json")

    args = parser.parse_args()
    if args.command == "export":
        export_docx(args.messages, args.docx)
    elif args.command == "sync":
        sync_from_docx(args.docx, args.messages)


if __name__ == "__main__":
    main()
